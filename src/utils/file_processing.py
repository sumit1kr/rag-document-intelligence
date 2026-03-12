import os
from datetime import datetime
from typing import List, Any, Dict, Union, Optional
import logging
import traceback
from PyPDF2 import PdfReader
import fitz  # PyMuPDF

from src.api.setup_api import logger

# -----------------------------
# File Processing
# -----------------------------
class FileProcessor:
    """Handles file upload and text extraction"""
    
    @staticmethod
    def extract_text(files: List[Any], use_ocr: bool = False) -> str:
        """Extract text from uploaded files"""
        if not files:
            return ""
        
        combined_texts = []
        processed_count = 0
        
        for file_obj in files:
            try:
                filename = getattr(file_obj, 'name', 'unknown')
                logger.info(f"Processing file: {filename}")
                
                # Determine file type
                file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
                
                content = ""
                
                # Process based on file type
                if file_extension == 'pdf':
                    content = FileProcessor._extract_pdf(file_obj)
                elif file_extension == 'txt':
                    content = FileProcessor._extract_text_file(file_obj)
                elif file_extension in ['docx', 'doc']:
                    content = FileProcessor._extract_word_document(file_obj, file_extension)
                elif file_extension in ['eml', 'msg']:
                    content = FileProcessor._extract_email(file_obj, file_extension)
                elif file_extension in ['png', 'jpg', 'jpeg'] and use_ocr:
                    content = FileProcessor._extract_image_ocr(file_obj)
                elif file_extension in ['png', 'jpg', 'jpeg'] and not use_ocr:
                    logger.info(f"Skipping image {filename} - OCR not enabled")
                    continue
                else:
                    logger.warning(f"Unsupported file type: {file_extension} for {filename}")
                    continue
                
                if content and content.strip():
                    combined_texts.append(f"=== Content from {filename} ===\n{content}")
                    processed_count += 1
                    logger.info(f"Successfully extracted text from {filename} ({len(content)} characters)")
                else:
                    logger.warning(f"No text extracted from {filename}")
                    
            except Exception as e:
                logger.error(f"Error processing file {getattr(file_obj, 'name', 'unknown')}: {str(e)}")
                logger.error(traceback.format_exc())
        
        result = "\n\n".join(combined_texts)
        logger.info(f"Processed {processed_count} files, total text length: {len(result)}")
        return result

    @staticmethod
    def _extract_pdf(file_obj: Any) -> str:
        """Extract text from PDF file"""
        try:
            # First try PyPDF2
            if hasattr(file_obj, 'read'):
                file_obj.seek(0)
            
            reader = PdfReader(file_obj)
            texts = []
            
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        texts.append(text)
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {e}")
            
            text = "\n".join(texts)
            
            # If PyPDF2 didn't work well, try PyMuPDF
            if not text.strip():
                logger.info("PyPDF2 extraction resulted in empty text, trying PyMuPDF")
                if hasattr(file_obj, 'read'):
                    file_obj.seek(0)
                    pdf_bytes = file_obj.read()
                else:
                    with open(file_obj, 'rb') as f:
                        pdf_bytes = f.read()
                
                doc = fitz.open(stream=pdf_bytes, filetype='pdf')
                texts = []
                for page in doc:
                    texts.append(page.get_text())
                doc.close()
                text = "\n".join(texts)
            
            return text
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ""

    @staticmethod
    def _extract_word_document(file_obj: Any, file_extension: str) -> str:
        """Extract text from Word documents (.docx and .doc)"""
        try:
            if file_extension == 'docx':
                return FileProcessor._extract_docx(file_obj)
            elif file_extension == 'doc':
                return FileProcessor._extract_doc(file_obj)
            else:
                logger.error(f"Unsupported Word document format: {file_extension}")
                return ""
                
        except Exception as e:
            logger.error(f"Word document extraction failed: {e}")
            return ""

    @staticmethod
    def _extract_docx(file_obj: Any) -> str:
        """Extract text from .docx files using python-docx"""
        try:
            from docx import Document
            
            if hasattr(file_obj, 'read'):
                file_obj.seek(0)
            
            doc = Document(file_obj)
            texts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    texts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    if row_texts:
                        texts.append(" | ".join(row_texts))
            
            return "\n".join(texts)
            
        except ImportError:
            logger.error("python-docx not installed. Please install it: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""

    @staticmethod
    def _extract_doc(file_obj: Any) -> str:
        """Extract text from .doc files using python-docx2txt"""
        try:
            import docx2txt
            
            if hasattr(file_obj, 'read'):
                file_obj.seek(0)
            
            # docx2txt can handle both .doc and .docx files
            text = docx2txt.process(file_obj)
            return text
            
        except ImportError:
            logger.error("docx2txt not installed. Please install it: pip install docx2txt")
            return ""
        except Exception as e:
            logger.error(f"DOC extraction failed: {e}")
            return ""

    @staticmethod
    def _extract_email(file_obj: Any, file_extension: str) -> str:
        """Extract text from email files (.eml, .msg)"""
        try:
            if file_extension == 'eml':
                return FileProcessor._extract_eml(file_obj)
            elif file_extension == 'msg':
                return FileProcessor._extract_msg(file_obj)
            else:
                logger.error(f"Unsupported email format: {file_extension}")
                return ""
                
        except Exception as e:
            logger.error(f"Email extraction failed: {e}")
            return ""

    @staticmethod
    def _extract_eml(file_obj: Any) -> str:
        """Extract text from .eml files"""
        try:
            import email
            from email import policy
            from bs4 import BeautifulSoup
            
            if hasattr(file_obj, 'read'):
                file_obj.seek(0)
                eml_content = file_obj.read()
            else:
                with open(file_obj, 'rb') as f:
                    eml_content = f.read()
            
            # Parse email
            msg = email.message_from_bytes(eml_content, policy=policy.default)
            
            email_text = []
            
            # Extract email headers
            email_text.append(f"From: {msg.get('from', 'Unknown')}")
            email_text.append(f"To: {msg.get('to', 'Unknown')}")
            email_text.append(f"Subject: {msg.get('subject', 'No Subject')}")
            email_text.append(f"Date: {msg.get('date', 'Unknown')}")
            email_text.append("---")
            
            # Extract email body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_maintype() == 'multipart':
                        continue
                    if part.get_content_maintype() == 'text':
                        content = part.get_content()
                        if part.get_content_subtype() == 'html':
                            # Convert HTML to text
                            soup = BeautifulSoup(content, 'html.parser')
                            content = soup.get_text()
                        email_text.append(content)
            else:
                content = msg.get_content()
                if msg.get_content_subtype() == 'html':
                    soup = BeautifulSoup(content, 'html.parser')
                    content = soup.get_text()
                email_text.append(content)
            
            return "\n".join(email_text)
            
        except ImportError as e:
            logger.error(f"Required library not installed: {e}")
            return ""
        except Exception as e:
            logger.error(f"EML extraction failed: {e}")
            return ""

    @staticmethod
    def _extract_msg(file_obj: Any) -> str:
        """Extract text from .msg files (Outlook message files)"""
        try:
            # For .msg files, we'll use a simple approach
            # In production, you might want to use libraries like extract-msg
            logger.warning(".msg file support is limited. Consider converting to .eml format.")
            
            if hasattr(file_obj, 'read'):
                file_obj.seek(0)
                content = file_obj.read()
            else:
                with open(file_obj, 'rb') as f:
                    content = f.read()
            
            # Try to extract readable text from binary content
            try:
                return content.decode('utf-8', errors='ignore')
            except:
                return content.decode('latin-1', errors='ignore')
                
        except Exception as e:
            logger.error(f"MSG extraction failed: {e}")
            return ""

    @staticmethod
    def _extract_text_file(file_obj: Any) -> str:
        """Extract text from plain text file"""
        try:
            if hasattr(file_obj, 'read'):
                content = file_obj.read()
            else:
                with open(file_obj, 'rb') as f:
                    content = f.read()
            
            # Try to decode as UTF-8, fallback to other encodings
            if isinstance(content, bytes):
                try:
                    return content.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        return content.decode('latin-1')
                    except UnicodeDecodeError:
                        return content.decode('utf-8', errors='ignore')
            
            return str(content)
            
        except Exception as e:
            logger.error(f"Text file extraction failed: {e}")
            return ""

    @staticmethod
    def _extract_image_ocr(file_obj: Any) -> str:
        """Extract text from image using OCR"""
        try:
            from PIL import Image
            import pytesseract
            
            if hasattr(file_obj, 'read'):
                file_obj.seek(0)
                img = Image.open(file_obj)
            else:
                img = Image.open(file_obj)
            
            text = pytesseract.image_to_string(img)
            return text
            
        except ImportError:
            logger.error("pytesseract not installed. Please install it for OCR support: pip install pytesseract")
            return ""
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return ""
