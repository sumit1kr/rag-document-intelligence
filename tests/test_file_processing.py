#!/usr/bin/env python3
"""
Test script for enhanced file processing functionality
Tests Word document and email file processing
"""

import os
import tempfile
from src.utils.file_processing import FileProcessor

def create_test_files():
    """Create test files for different formats"""
    test_files = []
    
    # Create a test .txt file
    txt_content = "This is a test text file.\nIt contains multiple lines.\nFor testing purposes."
    txt_file = tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False)
    txt_file.write(txt_content)
    txt_file.close()
    test_files.append(txt_file.name)
    
    # Create a test .docx file (simplified)
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("This is a test Word document.")
        doc.add_paragraph("It contains multiple paragraphs.")
        doc.add_paragraph("For testing the file processing system.")
        
        docx_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc.save(docx_file.name)
        docx_file.close()
        test_files.append(docx_file.name)
        print("✅ Created test .docx file")
    except ImportError:
        print("⚠️ python-docx not installed, skipping .docx test")
    
    # Create a test .eml file
    eml_content = """From: test@example.com
To: recipient@example.com
Subject: Test Email
Date: Mon, 1 Jan 2024 12:00:00 +0000
Content-Type: text/plain

This is a test email.
It contains email content for testing.
The email processing should extract this text.
"""
    eml_file = tempfile.NamedTemporaryFile(mode='w', suffix='.eml', delete=False)
    eml_file.write(eml_content)
    eml_file.close()
    test_files.append(eml_file.name)
    
    return test_files

def test_file_processing():
    """Test the file processing functionality"""
    print("🧪 Testing Enhanced File Processing")
    print("=" * 50)
    
    # Create test files
    test_files = create_test_files()
    
    # Test file processing
    for file_path in test_files:
        print(f"\n📁 Testing: {os.path.basename(file_path)}")
        
        # Create a file-like object for testing
        class FileObj:
            def __init__(self, path):
                self.name = path
                self._path = path
            
            def read(self):
                with open(self._path, 'rb') as f:
                    return f.read()
            
            def seek(self, pos):
                pass
        
        file_obj = FileObj(file_path)
        
        # Process the file
        try:
            content = FileProcessor.extract_text([file_obj])
            if content.strip():
                print(f"✅ Successfully extracted {len(content)} characters")
                print(f"📄 Preview: {content[:100]}...")
            else:
                print("❌ No content extracted")
        except Exception as e:
            print(f"❌ Error processing file: {e}")
    
    # Cleanup
    for file_path in test_files:
        try:
            os.unlink(file_path)
        except:
            pass
    
    print("\n" + "=" * 50)
    print("🎉 File processing test completed!")

if __name__ == "__main__":
    test_file_processing() 